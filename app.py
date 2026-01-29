"""
Comprehensive Pinellas County Property Lookup Test App
Tests: City auto-fill, Address, Property Use, Future Land Use, Zoning, Land Area
"""
import streamlit as st
import re
import json
import pathlib
import requests
from bs4 import BeautifulSoup
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

st.set_page_config(page_title="Pinellas Property Lookup Test", page_icon="🏠", layout="wide")

# ============================================================================
# PINELLAS CITY NAME MAPPING
# ============================================================================

# Map Pinellas County tax district codes to full city names
PINELLAS_CITY_MAP = {
    'SP': 'St. Petersburg',
    'ST PETERSBURG': 'St. Petersburg',
    'ST. PETERSBURG': 'St. Petersburg',
    'CLEARWATER': 'Clearwater',
    'CW': 'Clearwater',
    'CWD': 'Clearwater',
    'LARGO': 'Largo',
    'LA': 'Largo',
    'PINELLAS PARK': 'Pinellas Park',
    'PP': 'Pinellas Park',
    'DUNEDIN': 'Dunedin',
    'TARPON SPRINGS': 'Tarpon Springs',
    'TS': 'Tarpon Springs',
    'SEMINOLE': 'Seminole',
    'KENNETH CITY': 'Kenneth City',
    'GULFPORT': 'Gulfport',
    'MADEIRA BEACH': 'Madeira Beach',
    'MB': 'Madeira Beach',
    'REDINGTON BEACH': 'Redington Beach',
    'RB': 'Redington Beach',
    'TREASURE ISLAND': 'Treasure Island',
    'TI': 'Treasure Island',
    'ST PETE BEACH': 'St. Pete Beach',
    'SPB': 'St. Pete Beach',
    'SOUTH PASADENA': 'South Pasadena',
    'BELLEAIR': 'Belleair',
    'BELLEAIR BEACH': 'Belleair Beach',
    'BELLEAIR BLUFFS': 'Belleair Bluffs',
    'INDIAN ROCKS BEACH': 'Indian Rocks Beach',
    'IRB': 'Indian Rocks Beach',
    'INDIAN SHORES': 'Indian Shores',
    'NORTH REDINGTON BEACH': 'North Redington Beach',
    'NRB': 'North Redington Beach',
    'OLDSMAR': 'Oldsmar',
    'SAFETY HARBOR': 'Safety Harbor',
    # Unincorporated area codes
    'LFPW': 'Unincorporated Pinellas (Lealman)',
    'LEALMAN': 'Unincorporated Pinellas (Lealman)',
    'UNINCORPORATED': 'Unincorporated Pinellas',
    'COUNTY': 'Unincorporated Pinellas'
}

PINELLAS_CITY_LOOKUP_PATH = pathlib.Path(__file__).parent / "data" / "pinellas_county_cities_lookup.json"

def _load_pinellas_city_lookup():
    try:
        with open(PINELLAS_CITY_LOOKUP_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

PINELLAS_CITY_LOOKUP = _load_pinellas_city_lookup()
PINELLAS_CITY_NAMES = {
    name.strip().upper()
    for name, meta in PINELLAS_CITY_LOOKUP.items()
    if isinstance(meta, dict) and meta.get("type") == "city_app"
}

for city_name in PINELLAS_CITY_NAMES:
    PINELLAS_CITY_MAP.setdefault(city_name, city_name.title())

def expand_city_name(city_abbr):
    """Expand Pinellas city abbreviation to full name."""
    if not city_abbr:
        return 'Unincorporated Pinellas'
    
    city_upper = city_abbr.strip().upper()
    return PINELLAS_CITY_MAP.get(city_upper, city_abbr)

# ============================================================================
# ST. PETERSBURG ZONING/FLU LOOKUP TABLES
# ============================================================================

# St. Petersburg Future Land Use code to description mapping
FLU_DESCRIPTIONS = {
    'CBD': 'Central Business District',
    'CRD': 'Community Redevelopment District',
    'PR-R': 'Planned Redevelopment Residential',
    'PR-MU': 'Planned Redevelopment Mixed-Use',
    'PR-C': 'Planned Redevelopment Commercial',
    'RU': 'Residential Urban',
    'RL': 'Residential Low',
    'RLM': 'Residential Low Medium',
    'RM': 'Residential Medium',
    'RH': 'Residential High',
    'RVH': 'Residential Very High',
    'R/OL': 'Residential/Office Limited',
    'R/OG': 'Residential/Office General',
    'RFH': 'Resort Facilities High',
    'CG': 'Commercial General',
    'IL': 'Industrial Limited',
    'IG': 'Industrial General',
    'P': 'Preservation',
    'R/OS': 'Recreation/Open Space',
    'I': 'Institutional',
    'T/U': 'Transportation/Utility'
}

# St. Petersburg Zoning code to description mapping
ZONING_DESCRIPTIONS = {
    'NT-1': 'NEIGHBORHOOD TRADITIONAL SINGLE-FAMILY-1',
    'NT-2': 'NEIGHBORHOOD TRADITIONAL SINGLE-FAMILY-2',
    'NT-3': 'NEIGHBORHOOD TRADITIONAL SINGLE-FAMILY-3',
    'NT-4': 'NEIGHBORHOOD TRADITIONAL SINGLE-FAMILY-4',
    'NS-1': 'NEIGHBORHOOD SUBURBAN SINGLE-FAMILY-1',
    'NS-2': 'NEIGHBORHOOD SUBURBAN SINGLE-FAMILY-2',
    'NSM-1': 'NEIGHBORHOOD SUBURBAN MULTI-FAMILY-1',
    'NSM-2': 'NEIGHBORHOOD SUBURBAN MULTI-FAMILY-2',
    'NPUD-1': 'NEIGHBORHOOD PLANNED UNIT DEVELOPMENT-1',
    'NPUD-2': 'NEIGHBORHOOD PLANNED UNIT DEVELOPMENT-2',
    'NPUD-3': 'NEIGHBORHOOD PLANNED UNIT DEVELOPMENT-3',
    'NTM-1': 'NEIGHBORHOOD TRADITIONAL MIXED-RESIDENTIAL-1',
    'NMH': 'NEIGHBORHOOD SUBURBAN MOBILE HOME',
    'NSE': 'NEIGHBORHOOD SUBURBAN SINGLE-FAMILY',
    'DC-C': 'DOWNTOWN CENTER-CORE',
    'DC-1': 'DOWNTOWN CENTER-1',
    'DC-2': 'DOWNTOWN CENTER-2',
    'DC-3': 'DOWNTOWN CENTER-3',
    'DC-P': 'DOWNTOWN CENTER-PRESERVATION',
    'CRT-1': 'CORRIDOR RESIDENTIAL TRADITIONAL-1',
    'CRT-2': 'CORRIDOR RESIDENTIAL TRADITIONAL-2',
    'CRS-1': 'CORRIDOR RESIDENTIAL SUBURBAN-1',
    'CRS-2': 'CORRIDOR RESIDENTIAL SUBURBAN-2',
    'CCT-1': 'CORRIDOR COMMERCIAL TRADITIONAL-1',
    'CCT-2': 'CORRIDOR COMMERCIAL TRADITIONAL-2',
    'CCS-1': 'CORRIDOR COMMERCIAL SUBURBAN-1',
    'CCS-2': 'CORRIDOR COMMERCIAL SUBURBAN-2',
    'RC-1': 'RETAIL CENTER-1',
    'RC-2': 'RETAIL CENTER-2',
    'RC-3': 'RETAIL CENTER-3',
    'EC-1': 'EMPLOYMENT CENTERS-1',
    'EC-2': 'EMPLOYMENT CENTERS-2',
    'IS': 'INDUSTRIAL SUBURBAN',
    'IT': 'INDUSTRIAL TRADITIONAL',
    'IC': 'INSTITUTIONAL CENTER',
    'P': 'PRESERVATION',
    'WATER': 'WATER'
}

# NOTE: Unincorporated Pinellas zoning and FLU coded values are now fetched 
# dynamically from layer metadata using fetch_coded_values() function.
# Hardcoded fallbacks kept below in case metadata fetch fails.

# Unincorporated Pinellas County Zoning code to description mapping (FALLBACK)
UNINCORPORATED_ZONING_DESCRIPTIONS = {
    'AL': 'Aquatic Lands',
    'AL-CO': 'Aquatic Lands - Conditional Overlay',
    'AL-W': 'Aquatic Lands - Wellhead Protection Overlay',
    'AL-W-CO': 'Aquatic Lands - Wellhead Protection Overlay - Conditional Overlay',
    'C-1': 'Neighborhood Commercial',
    'C-1-CO': 'Neighborhood Commercial - Conditional Overlay',
    'C-1-H': 'Neighborhood Commerical - Historic District',
    'C-1-W': 'Neighborhood Commercial - Wellhead Protection Overlay',
    'C-1-W-CO': 'Neighborhood Commercial - Wellhead Protection Overlay - Conditional Overlay',
    'CP': 'Commercial Parkway',
    'CP-CO': 'Commercial Parkway - Conditional Overlay',
    'CP-W': 'Commercial Parkway - Wellhead Protection Overlay',
    'CP-W-CO': 'Commercial Parkway - Wellhead Protection Overlay - Conditional Overlay',
    'CR': 'Commercial Recreation',
    'CR-CO': 'Commercial Recreation - Conditional Overlay',
    'CR-W': 'Commercial Recreation - Wellhead Protection Overlay',
    'CR-W-CO': 'Commercial Recreation - Wellhead Protection Overlay - Conditional Overlay',
    'E-1': 'Employment 1',
    'E-1-CO': 'Employment 1 - Conditional Overlay',
    'E-1-W': 'Employment 1 - Wellhead Protection Overlay',
    'E-1-W-CO': 'Employment 1 - Wellhead Protection Overlay - Conditional Overlay',
    'E-2': 'Employment 2',
    'E-2-CO': 'Employment 2 - Conditional Overlay',
    'E-2-W': 'Employment 2 - Wellhead Protection Overlay',
    'E-2-W-CO': 'Employment 2 - Wellhead Protection Overlay - Conditional Overlay',
    'FBC': 'Form Based Code District',
    'FBC-CO': 'Form Based Code District - Conditional Overlay',
    'FBC-W': 'Form Based Code District - Wellhead Protection Overlay',
    'FBC-W-CO': 'Form Based Code District - Wellhead Protection Overlay - Conditional Overlay',
    'FBR': 'Facilities-Based Recreation',
    'FBR-CO': 'Facilities-Based Recreation - Conditional Overlay',
    'FBR-W': 'Facilities-Based Recreation - Wellhead Protection Overlay',
    'FBR-W-CO': 'Facilities-Based Recreation - Wellhead Protection Overlay - Conditional Overlay',
    'GI': 'General Institutional',
    'GI-CO': 'General Institutional - Conditional Overlay',
    'GI-W': 'General Institutional - Wellhead Protection Overlay',
    'GI-W-CO': 'General Institutional - Wellhead Protection Overlay - Conditional Overlay',
    'GO': 'General Office',
    'GO-CO': 'General Office - Conditional Overlay',
    'GO-W': 'General Office - Wellhead Protection Overlay',
    'GO-W-CO': 'General Office - Wellhead Protection Overlay - Conditional Overlay',
    'I': 'Industrial',
    'I-CO': 'Industrial - Conditional Overlay',
    'I-W': 'Industrial - Wellhead Protection Overlay',
    'I-W-CO': 'Industrial - Wellhead Protection Overlay',
    'IPD': 'Industrial Planned Development',
    'IPD-CO': 'Industrial Planned Development - Conditional Overlay',
    'IPD-W': 'Industrial Planned Development - Wellhead Protection Overlay',
    'IPD-W-CO': 'Industrial Planned Development - Wellhead Protection Overlay - Conditional Overlay',
    'LI': 'Limited Institutional',
    'LI-CO': 'Limited Institutional - Conditional Overlay',
    'LI-W': 'Limited Institutional - Wellhead Protection Overlay',
    'LI-W-CO': 'Limited Institutional - Wellhead Protection Overlay - Conditional Overlay',
    'LO': 'Limited Office',
    'LO-CO': 'Limited Office - Conditional Overlay',
    'LO-W': 'Limited Office - Wellhead Protection Overlay',
    'LO-W-CO': 'Limited Office - Wellhead Protection Overlay - Conditional Overlay',
    'MXD': 'Mixed-Use District',
    'MXD-CO': 'Mixed-Use District - Conditional Overlay',
    'MXD-W': 'Mixed-Use District - Wellhead Protection Overlay',
    'MXD-W-CO': 'Mixed-Use District - Wellhead Protection Overlay - Conditional Overlay',
    'OPH-D': 'Old Palm Harbor Downtown',
    'OPH-D-CO': 'Old Palm Harbor Downtown - Conditional Overlay',
    'OPH-D-H': 'Old Palm Harbor Downtown - Historic District',
    'OPH-D-W': 'Old Palm Harbor Downtown - Wellhead Protection Overlay',
    'OPH-D-W-CO': 'Old Palm Harbor Downtown - Wellhead Protection Overlay - Conditional Overlay',
    'P-C': 'Preservation Conservation',
    'P-C-CO': 'Preservation Conservation - Conditional Overlay',
    'P-C-W': 'Preservation Conservation - Wellhead Protection Overlay',
    'P-C-W-CO': 'Preservation Conservation - Wellhead Protection Overlay - Conditional Overlay',
    'P-RM': 'Preservation Resource Management',
    'P-RM-CO': 'Preservation Resource Management - Conditional Overlay',
    'P-RM-W': 'Preservation Resource Management - Wellhead Protection Overlay',
    'P-RM-W-CO': 'Preservation Resource Management - Wellhead Protection Overlay - Conditional Overlay',
    'P.C.AIRPORT': 'PC Airport',
    'P.C.AIRPORT-CO': 'PC Airport - Conditional Overlay',
    'P.C.AIRPORT-W': 'PC Airport - Wellhead Protection Overlay',
    'P.C.AIRPORT-W-CO': 'PC Airport- Wellhead Protection Overlay - Conditional Overlay',
    'R-1': 'Single Family Residential (9,500 SF Min)',
    'R-1-CO': 'Single Family Residential (9,500 SF Min) - Conditional Overlay',
    'R-1-W': 'Single Family Residential (9,500 SF Min) - Wellhead Protection Overlay',
    'R-1-W-CO': 'Single Family Residential (9,500 SF Min) - Wellhead Protection Overlay - Conditional Overlay',
    'R-2': 'Single Family Residential (7,500 SF Min)',
    'R-2-CO': 'Single Family Residential (7,500 SF Min) - Conditional Overlay',
    'R-2-W': 'Single Family Residential (7,500 SF Min) - Wellhead Protection Overlay',
    'R-2-W-CO': 'Single Family Residential (7,500 SF Min) - Wellhead Protection Overlay - Conditional Overlay',
    'R-3': 'Single Family Residential (6,000 SF Min)',
    'R-3-CO': 'Single Family Residential (6,000 SF Min) - Conditional Overlay',
    'R-3-H': 'Single Family Residential (6,000 SF Min) - Historic District',
    'R-3-W': 'Single Family Residential (6,000 SF Min) - Wellhead Protection Overlay',
    'R-3-W-CO': 'Single Family Residential (6,000 SF Min) - Wellhead Protection Overlay - Conditional Overlay',
    'R-4': 'One, Two and Three Family Residential',
    'R-4-CO': 'One, Two and Three Family Residential - Conditional Overlay',
    'R-4-W': 'One, Two and Three Family Residential - Wellhead Protection Overlay',
    'R-4-W-CO': 'One, Two and Three Family Residential - Wellhead Protection Overlay - Conditional Overlay',
    'R-5': 'Urban Residential District',
    'R-5-CO': 'Urban Residential District - Conditional Overlay',
    'R-5-W': 'Urban Residential District - Wellhead Protection Overlay',
    'R-5-W-CO': 'Urban Residential District - Wellhead Protection Overlay - Conditional Overlay',
    'R-A': 'Residential Agriculture',
    'R-A-CO': 'Residential Agriculture - Conditional Overlay',
    'R-A-W': 'Residential Agriculture - Wellhead Protection Overlay',
    'R-A-W-CO': 'Residential Agriculture - Wellhead Protection Overlay - Conditional Overlay',
    'R-E': 'Residential Estate',
    'R-E-C-T': 'Residential Estate - Transient Accommodation Overlay',
    'R-E-CO': 'Residential Estate - Conditional Overlay',
    'R-E-W': 'Residential Estate - Wellhead Protection Overlay',
    'R-E-W-CO': 'Residential Estate - Wellhead Protection Overlay - Conditional Overlay',
    'R-R': 'Rural Residential',
    'R-R-CO': 'Rural Residential - Conditional Overlay',
    'R-R-H': 'Rural Residential - Historic District',
    'R-R-W': 'Rural Residnetial - Wellhead Protection Overlay',
    'R-R-W-CO': 'Rural Residential - Wellhead Protection Overlay - Conditional Overlay',
    'RBR': 'Resource-Based Recreation',
    'RBR-CO': 'Resource-Based Recreation - Conditional Overlay',
    'RBR-W': 'Resource-Based Recreation - Wellhead Protection Overlay',
    'RBR-W-CO': 'Resource-Based Recreation - Wellhead Protection Overlay - Conditional Overlay',
    'RM': 'Multi-Family Residential (see FLUM for density)',
    'RM-CO': 'Multi-Family Residential (see FLUM for density) - Conditional Overlay',
    'RM-W': 'Multi-Family Residential (see FLUM for density) - Wellhead Protection Overlay',
    'RM-W-CO': 'Multi-Family Residential (see FLUM for density) - Wellhead Protection Overlay - Conditional Overlay',
    'RMH': 'Residential Mobile/Manufactured Home',
    'RMH-CO': 'Residential Mobile/Manufactured Home - Conditional Overlay',
    'RMH-W': 'Residential Mobile/Manufactured Home - Wellhead Protection Overlay',
    'RMH-W-CO': 'Residential Mobile/Manufactured Home - Wellhead Protection Overlay - Conditional Overlay',
    'RPD': 'Residential Planned Development (see FLUM for density)',
    'RPD-CO': 'Residential Planned Development (see FLUM for density) - Conditional Overlay',
    'RPD-W': 'Residential Planned Developlment (see FLUM for density) - Wellhead Protection Overlay',
    'RPD-W-CO': 'Residential Planned Development (see FLUM for density) - Wellhead Protection Overlay - Conditional Overlay',
    'UZ': 'Unknown Zoning',
    'UZ-CO': 'Unknown Zoning - Conditional Overlay',
    'UZ-W': 'Unknown Zoning - Wellhead Protection Overlay',
    'UZ-W-CO': 'Unknown Zoning - Wellhead Protection Overlay - Conditional Overlay',
    'OPH-D-W': 'Old Palm Harbor Downtown - Wellhead Protection Overlay',
    'C-T': 'Transient Accommodation Overlay',
    'HPO': 'Historic Preservation Overlay',
    'E-1-C-T': 'Employment 1 - Transient Accommodation Overlay',
    'C-2': 'General Commercial and Services',
    'DPH-FBC': 'Downtown Palm Harbor Form Based Code',
    'C-2-C-T': 'General Commercial and Services Transient Accommodations Overlay',
    'L-FBC': 'Lealman - Form Based Code',
    'C-2-CO': 'General Commercial and Services - Conditional Overlay',
    'C-2-H': 'General Commercial and Services - Historic District',
    'C-2-W': 'General Commercial and Services - Wellhead Protection Overlay',
    'C-2-W-CO': 'General Commercial and Services - Wellhead Proteciton Overlay - Conditional Overlay',
}

# Unincorporated Pinellas County Future Land Use code to description mapping (FALLBACK)
UNINCORPORATED_FLU_DESCRIPTIONS = {
    'RR': 'Residential Rural',
    'RE': 'Residential Estate',
    'RS': 'Residential Suburban',
    'RL': 'Residential Low',
    'RU': 'Residential Urban',
    'RLM': 'Residential Low Medium',
    'RM': 'Residential Medium',
    'RH': 'Residential High',
    'PR-I': 'Planned Redevelopment - Industrial',
    'RFO': 'Resort Facilities',
    'PR-C': 'Planned Redevelopment - Commercial',
    'NO-DES': 'No Designation',
    'CN': 'Commercial Neighborhood',
    'CG': 'Commercial General',
    'CR': 'Commercial Recreation',
    'IL': 'Industrial Limited',
    'IG': 'Industrial General',
    'P': 'Preservation',
    'PR-MU': 'Planned Redevelopment - Mixed Use',
    'I': 'Institutional',
    'PR-R': 'Planned Redevelopment - Residential',
    'TU': 'Transportation/Utilities',
    'ROR': 'Residential/Office/Retail',
    'ROL': 'Residential/Office/Limited',
    'ROG': 'Residential/Office/General',
    'ROS': 'Recreation/Open Space',
    'PSP': 'Public/Semi-Public',
    'RVH': 'Residential Very High',
    'RFM': 'Resort Facilities Medium',
    'RFH': 'Resort Facilities High',
    'CRD': 'Community Redevelopment Dist',
    'CBD': 'Central Business District',
    'CL': 'Commercial Limited',
    'RFO-P': 'Resort Facilities Overlay/Perm',
    'RFO-T': 'Resort Facilities Overlay/Temp',
    'WDF': 'Water Drainage Feature',
    'WF': 'Water Feature',
    'WATER': 'WATER',
    'ROAD': 'ROAD',
    'P-RM': 'Preservation - Resource Management',
    'MUNI': 'MUNICIPALITY',
    'NO-D-W': 'No Designation Uninc Water',
    'MUNI-W': 'Municipal Open Water',
    'CRD-AC': 'Community Redevelop-Activity Ctr',
    'AC': 'AC',
    'AC-P': 'AC-P',
    'RM-12.5': 'RM-12.5',
    'TU-O': 'Transportation/Utility Overlay',
    'E': 'Employment',
    'AC-N': 'Activity Center - Neighborhood',
    'AC-C': 'Activity Center - Community',
    'AC-M': 'Activity Center - Major',
    'MUC-P': 'Mixed Use Corridor - Primary',
    'MUC-S': 'Mixed Use Corridor - Secondary',
    'MUC-P-C': 'Mixed Use Corridor - Primary - Commerce',
    'MUC-SU-NP': 'Mixed Use Corridor - Supporting - Neighborhood Park',
    'MUC-SU-LT': 'Mixed Use Corridor - Supporting - Local Trade',
}

# ============================================================================
# AUTOMATED CODED VALUE EXTRACTION FROM ARCGIS LAYERS
# ============================================================================

@st.cache_data(ttl=86400, show_spinner=False)
def fetch_coded_values(layer_url, field_name):
    """
    Automatically fetch coded value domain from an ArcGIS layer.
    
    Args:
        layer_url: Base layer URL (e.g., "https://...MapServer/1")
        field_name: Field with coded values (e.g., "ZONEDESC")
    
    Returns:
        dict: {code: description} mapping, or empty dict if not found
    """
    session = get_resilient_session()
    
    try:
        # Fetch layer metadata
        metadata_url = f"{layer_url}?f=json"
        response = session.get(metadata_url, timeout=15)
        response.raise_for_status()
        metadata = response.json()
        
        # Find the field with coded values
        for field in metadata.get('fields', []):
            if field.get('name') == field_name:
                domain = field.get('domain')
                if domain and domain.get('type') == 'codedValue':
                    # Extract code -> name mappings
                    coded_values = domain.get('codedValues', [])
                    return {item['code']: item['name'] for item in coded_values}
        
        return {}
    
    except Exception as e:
        # Return empty dict on error - calling code should handle
        return {}

# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def validate_parcel_id(parcel_id: str):
    """Validate parcel/folio ID input."""
    if not parcel_id:
        return False, "Parcel ID cannot be empty"
    
    if len(parcel_id) > 30:
        return False, "Parcel ID must be 30 characters or less"
    
    # Allowlist: only alphanumeric, dashes, spaces, periods
    if not re.match(r'^[A-Za-z0-9\-\s\.]+$', parcel_id):
        return False, "Invalid characters in parcel ID"
    
    return True, ""

def sanitize_for_sql(value: str) -> str:
    """Sanitize string for use in SQL WHERE clause."""
    return value.strip().replace("'", "''")

def strip_dor_code(land_use_text):
    """Remove Florida DOR code prefix from land use descriptions."""
    if not land_use_text:
        return ''
    
    text = land_use_text.strip()
    
    if text and text[0].isdigit():
        parts = text.split(' ', 1)
        if len(parts) > 1:
            return parts[1].strip()
    
    return text

# ============================================================================
# HTTP SESSION WITH RETRY LOGIC
# ============================================================================

@st.cache_resource
def get_resilient_session():
    """Create HTTP session with automatic retry logic."""
    session = requests.Session()
    
    retry_strategy = Retry(
        total=3,
        backoff_factor=1.0,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["GET", "POST"]
    )
    
    adapter = HTTPAdapter(max_retries=retry_strategy)
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    
    return session



# ============================================================================
# PINELLAS GIS (PARCEL GEOMETRY -> ZONING/FLU) HELPERS
# ============================================================================
PINELLAS_GIS_BASE_URL = "https://gis.pinellascounty.org/arcgis/rest/services"
PINELLAS_GIS_SERVICES = {
    "parcels_layer": "Parcel/Parcels/MapServer/0",
    "zoning_layer": "Planning/ZoningMap/MapServer/0",
    "flum_layer": "Planning/FLUM/MapServer/0",
}

def _arcgis_get_json(session, url, params, timeout=20):
    """GET JSON from ArcGIS REST; raise on HTTP errors; return dict."""
    resp = session.get(url, params={**params, "f": "json"}, timeout=timeout)
    resp.raise_for_status()
    data = resp.json()
    if isinstance(data, dict) and data.get("error"):
        # Surface server error message
        msg = data["error"].get("message", "ArcGIS REST error")
        details = data["error"].get("details") or []
        raise RuntimeError(f"{msg} {' | '.join(map(str, details))}".strip())
    return data

def _pinellas_query_layer(layer_path, params, timeout=20):
    session = get_resilient_session()
    url = f"{PINELLAS_GIS_BASE_URL}/{layer_path}/query"
    return _arcgis_get_json(session, url, params, timeout=timeout)

def _pinellas_where_variants(parcel_field, parcel_id):
    """
    Build a list of WHERE clause variants for common parcel-id formatting cases.
    Tries:
      - original value (quoted)
      - digits-only value (quoted and unquoted if numeric)
    """
    pid = str(parcel_id).strip()
    pid_digits = "".join(ch for ch in pid if ch.isdigit())
    values = [pid]
    if pid_digits and pid_digits != pid:
        values.append(pid_digits)

    wheres = []
    for v in values:
        safe = v.replace("'", "''")
        # Text field form
        wheres.append(f"{parcel_field} = '{safe}'")
        # Numeric field form (only if digits)
        if v.isdigit():
            wheres.append(f"{parcel_field} = {v}")
    return wheres

def pinellas_get_parcel_feature(parcel_id):
    """Return first parcel feature (attributes + geometry) from Pinellas parcels service."""
    candidate_fields = ["PARCELID", "PARCEL_ID", "PIN", "APN", "FOLIO", "FOLIO_NUMBER", "PARCELNO", "PARCEL_NO"]

    last_error = None
    for fld in candidate_fields:
        for where in _pinellas_where_variants(fld, parcel_id):
            try:
                data = _pinellas_query_layer(
                    PINELLAS_GIS_SERVICES["parcels_layer"],
                    {
                        "where": where,
                        "outFields": "*",
                        "returnGeometry": "true",
                        # Ask for WGS84, but do not assume it will be honored.
                        "outSR": "4326",
                        "resultRecordCount": 1,
                    },
                )
                feats = data.get("features") or []
                if feats:
                    return {"success": True, "field": fld, "where": where, "feature": feats[0]}
            except Exception as e:
                last_error = str(e)
                continue

    return {"success": False, "error": "Parcel geometry not found in parcels service", "detail": last_error}

def _polygon_centroid_xy(rings):
    """Compute an area-weighted centroid for the first ring. Falls back to bbox center."""
    if not rings or not isinstance(rings, list) or not rings[0]:
        return None
    ring = rings[0]
    # Ensure closed ring
    if ring[0] != ring[-1]:
        ring = ring + [ring[0]]
    # Shoelace centroid (planar). Works adequately for small polygons.
    area2 = 0.0
    cx6 = 0.0
    cy6 = 0.0
    for i in range(len(ring) - 1):
        x0, y0 = ring[i]
        x1, y1 = ring[i + 1]
        cross = x0 * y1 - x1 * y0
        area2 += cross
        cx6 += (x0 + x1) * cross
        cy6 += (y0 + y1) * cross
    if abs(area2) > 1e-12:
        cx = cx6 / (3.0 * area2)
        cy = cy6 / (3.0 * area2)
        return float(cx), float(cy)
    # Fallback: bbox center
    xs = [pt[0] for pt in ring if isinstance(pt, (list, tuple)) and len(pt) >= 2]
    ys = [pt[1] for pt in ring if isinstance(pt, (list, tuple)) and len(pt) >= 2]
    if not xs or not ys:
        return None
    return float((min(xs) + max(xs)) / 2.0), float((min(ys) + max(ys)) / 2.0)

def _pick_attr(attrs, keys):
    for k in keys:
        if k in attrs and attrs.get(k) not in (None, ""):
            return attrs.get(k)
    # Case-insensitive fallback
    lower = {str(k).lower(): k for k in attrs.keys()}
    for k in keys:
        lk = str(k).lower()
        if lk in lower:
            v = attrs.get(lower[lk])
            if v not in (None, ""):
                return v
    return None

def pinellas_lookup_zoning_flu_via_parcel_geometry(parcel_id):
    """
    Deterministic zoning + FLU lookup:
      parcel_id -> parcel polygon geometry -> zoning/flum spatial query (polygon intersects).

    This avoids centroid edge cases and SR mismatches (we pass the parcel geometry + inSR).
    """
    session = get_resilient_session()
    try:
        parcel_res = pinellas_get_parcel_feature(parcel_id)
        if not parcel_res.get("success"):
            return {"success": False, "error": parcel_res.get("error", "Parcel geometry lookup failed"), "_debug": parcel_res}

        feature = parcel_res["feature"]
        geom = (feature or {}).get("geometry") or {}
        if not geom:
            return {"success": False, "error": "Parcel feature returned without geometry", "_debug": parcel_res}

        # Determine spatial reference of returned geometry
        sr = (geom.get("spatialReference") or {}).get("wkid") or 4326

        def _poly_query(layer_path):
            url = f"{PINELLAS_GIS_BASE_URL}/{layer_path}/query"
            params = {
                "where": "1=1",
                "geometry": json.dumps(geom),
                "geometryType": "esriGeometryPolygon",
                "inSR": str(sr),
                "spatialRel": "esriSpatialRelIntersects",
                "outFields": "*",
                "returnGeometry": "false",
                "resultRecordCount": 5,
            }
            return _arcgis_get_json(session, url, params)

        zoning_data = _poly_query(PINELLAS_GIS_SERVICES["zoning_layer"])
        z_feats = zoning_data.get("features") or []
        z_attrs = (z_feats[0] or {}).get("attributes") if z_feats else {}

        zoning_code = _pick_attr(z_attrs or {}, [
            "ZONING", "Zoning", "ZONING_CLASSIFICATION", "ZONING_CLASS", "ZONINGCODE",
            "DISTRICT", "ZONE", "CLASSNAME", "ClassName"
        ])
        zoning_desc = _pick_attr(z_attrs or {}, [
            "ZONING_DESC", "ZONING_DESCRIPTION", "ZoningDesc", "Description",
            "ZONEDESC", "ZONINGDESC", "DISTRICT_DESC", "ZONE_DESC"
        ])

        flum_data = _poly_query(PINELLAS_GIS_SERVICES["flum_layer"])
        f_feats = flum_data.get("features") or []
        f_attrs = (f_feats[0] or {}).get("attributes") if f_feats else {}

        flu_code = _pick_attr(f_attrs or {}, [
            "FLUM", "FLU", "FUTURELANDUSE", "FUTURE_LAND_USE",
            "CATEGORY", "CLASSNAME", "ClassName"
        ])
        flu_desc = _pick_attr(f_attrs or {}, [
            "FLUM_DESC", "FLU_DESC", "DESCRIPTION", "FutureLandUseDesc", "DESC"
        ])

        debug = {
            "parcel_field": parcel_res.get("field"),
            "parcel_where": parcel_res.get("where"),
            "parcel_sr": sr,
            "zoning_feature_count": len(z_feats),
            "flum_feature_count": len(f_feats),
        }

        if not any([zoning_code, zoning_desc, flu_code, flu_desc]):
            return {"success": False, "error": "No zoning/FLU returned from GIS layers", "_debug": debug}

        return {
            "success": True,
            "zoning_code": str(zoning_code) if zoning_code is not None else "",
            "zoning_description": str(zoning_desc) if zoning_desc is not None else "",
            "future_land_use": str(flu_code) if flu_code is not None else "",
            "future_land_use_description": str(flu_desc) if flu_desc is not None else "",
            "_gis_source": "pinellascounty.org ArcGIS (parcel polygon)",
            "_debug": debug,
        }

    except Exception as e:
        return {"success": False, "error": f"GIS zoning/FLU lookup error: {str(e)}"}


# ============================================================================
# PCPAO API LOOKUP (THE KEY FUNCTION!)
# ============================================================================

def scrape_pinellas_property(parcel_id):
    """
    Query Pinellas County Property Appraiser searchProperty API.
    This is the backend API that the PCPAO website uses.
    """
    session = get_resilient_session()
    
    url = "https://www.pcpao.gov/dal/quicksearch/searchProperty"
    
    # Normalize Pinellas parcel ID format
    normalized_parcel = parcel_id.strip()
    
    # If no dashes and 18 digits, add dashes in Pinellas format
    if '-' not in normalized_parcel and len(normalized_parcel) == 18:
        # Format: XX-XX-XX-XXXXX-XXX-XXXX
        normalized_parcel = f"{normalized_parcel[0:2]}-{normalized_parcel[2:4]}-{normalized_parcel[4:6]}-{normalized_parcel[6:11]}-{normalized_parcel[11:14]}-{normalized_parcel[14:18]}"
    
    # Build the POST data - mimics the DataTables request format
    payload = {
        'draw': '1',
        'start': '0',
        'length': '10',
        'search[value]': '',
        'search[regex]': 'false',
        'input': normalized_parcel,
        'searchsort': 'parcel_number',
        'url': 'https://www.pcpao.gov'
    }
    
    # Add column definitions (required by DataTables API)
    for i in range(11):
        payload[f'columns[{i}][data]'] = str(i)
        payload[f'columns[{i}][name]'] = ''
        payload[f'columns[{i}][searchable]'] = 'true'
        payload[f'columns[{i}][orderable]'] = 'true' if i >= 2 else 'false'
        payload[f'columns[{i}][search][value]'] = ''
        payload[f'columns[{i}][search][regex]'] = 'false'
    
    try:
        response = session.post(url, data=payload, timeout=15)
        response.raise_for_status()
        
        data = response.json()
        
        # Check if we got results
        if data.get('recordsTotal', 0) == 0:
            return {'success': False, 'error': 'Parcel not found in PCPAO database'}
        
        # Parse the HTML response within the JSON
        if not data.get('data') or len(data['data']) == 0:
            return {'success': False, 'error': 'No property data returned'}
        
        # Get first result
        result_row = data['data'][0]
        
        # Extract data from HTML snippets
        # Column 2: Owner name
        owner_html = result_row[2] if len(result_row) > 2 else ''
        owner_soup = BeautifulSoup(owner_html, 'lxml')
        owner = owner_soup.get_text(strip=True)
        
        # Column 5: Address
        address_html = result_row[5] if len(result_row) > 5 else ''
        address_soup = BeautifulSoup(address_html, 'lxml')
        address = address_soup.get_text(strip=True)
        
        # Column 6: Current Tax District (this IS the city, but may be abbreviated)
        tax_dist_html = result_row[6] if len(result_row) > 6 else ''
        tax_dist_soup = BeautifulSoup(tax_dist_html, 'lxml')
        tax_district = tax_dist_soup.get_text(strip=True)
        
        # Expand abbreviated city name (e.g., "SP" -> "St. Petersburg")
        city = expand_city_name(tax_district)
        
        # Column 7: Property Use / DOR Code
        use_html = result_row[7] if len(result_row) > 7 else ''
        use_soup = BeautifulSoup(use_html, 'lxml')
        property_use = use_soup.get_text(strip=True)
        
        # Column 8: Legal Description
        legal_html = result_row[8] if len(result_row) > 8 else ''
        legal_soup = BeautifulSoup(legal_html, 'lxml')
        legal_desc = legal_soup.get_text(strip=True)
        
        # Get acreage from detail page
        sqft = None
        acres = None
        zip_code = None
        
        try:
            # Strap transformation: swap first and third segments
            parts = normalized_parcel.split('-')
            if len(parts) == 6:
                parts[0], parts[2] = parts[2], parts[0]
                strap = ''.join(parts)
            else:
                strap = normalized_parcel.replace('-', '')
            
            # Build detail URL
            detail_url = (
                f"https://www.pcpao.gov/property-details?"
                f"s={strap}&"
                f"input={normalized_parcel}&"
                f"search_option=parcel_number"
            )
            
            # Fetch and parse
            html = session.get(detail_url, timeout=30).text
            soup = BeautifulSoup(html, "html.parser")
            text = soup.get_text(" ", strip=True)
            
            # Match pattern: "Land Area: ≅ 59,560 sf | ≅ 1.36 acres"
            m = re.search(r"Land Area:\s*≅\s*([\d,]+)\s*sf\s*\|\s*≅\s*([\d.]+)\s*acres", text)
            if m:
                sqft = int(m.group(1).replace(",", ""))
                acres = float(m.group(2))
            
            # Extract ZIP code from detail page (format: "FL 33703" or "FL33703")
            zip_match = re.search(r'FL\s*(\d{5})', text)
            if zip_match:
                zip_code = zip_match.group(1)
        except Exception:
            pass  # If detail page fails, sqft, acres, and zip_code remain None
        
        return {
            'success': True,
            'address': address,
            'city': city,
            'zip': zip_code or '',
            'owner': owner,
            'land_use': strip_dor_code(property_use),
            'zoning': 'Contact City/County for zoning info',
            'site_area_sqft': f"{sqft:,}" if sqft else None,
            'site_area_acres': f"{acres:.2f}" if acres else None,
            'legal_description': legal_desc,
            'error': None
        }
    
    except Exception as e:
        return {'success': False, 'error': f'Error querying PCPAO API: {str(e)}'}

# ============================================================================
# ST. PETERSBURG ZONING LAYER LOOKUP
# ============================================================================

def is_unincorporated(city_name):
    """Check if city is unincorporated Pinellas."""
    if not city_name:
        return True

    city_upper = city_name.upper()
    if city_upper in PINELLAS_CITY_MAP:
        return "UNINCORPORATED" in PINELLAS_CITY_MAP[city_upper].upper()
    if city_upper in PINELLAS_CITY_NAMES:
        return False
    
    unincorporated_indicators = [
        'UNINCORPORATED',
        'LFPW',  # Lealman area
        'LEALMAN',
        'COUNTY',
        'PINELLAS COUNTY'
    ]
    
    return any(indicator in city_upper for indicator in unincorporated_indicators)

def lookup_unincorporated_zoning(address):
    """
    Lookup zoning and FLU for unincorporated Pinellas County areas.
    Uses PublicWebGIS/Landuse_Zoning/MapServer with automated coded value extraction.
    
    Returns: dict with zoning_code, zoning_description, future_land_use, future_land_use_description
    """
    if not address:
        return {'success': False, 'error': 'Address required for zoning lookup'}
    

    # FAST PATH (deterministic): If we have a parcel_id, try parcel-geometry-based GIS lookup first.
    if parcel_data and parcel_data.get('parcel_id'):
        gis_res = pinellas_lookup_zoning_flu_via_parcel_geometry(parcel_data.get('parcel_id'))
        if gis_res.get('success'):
            return gis_res

    session = get_resilient_session()
    
    try:
        # Fetch coded value domains dynamically (cached for 24 hours)
        zoning_lookup = fetch_coded_values(
            "https://egis.pinellas.gov/gis/rest/services/PublicWebGIS/Landuse_Zoning/MapServer/1",
            "ZONEDESC"
        )
        flu_lookup = fetch_coded_values(
            "https://egis.pinellas.gov/gis/rest/services/PublicWebGIS/Landuse_Zoning/MapServer/0",
            "LANDUSEDESC"
        )
        
        # Fallback to hardcoded if fetch failed
        if not zoning_lookup:
            zoning_lookup = UNINCORPORATED_ZONING_DESCRIPTIONS
        if not flu_lookup:
            flu_lookup = UNINCORPORATED_FLU_DESCRIPTIONS
        
        # Step 1: Geocode the address
        search_url = "https://geocode.arcgis.com/arcgis/rest/services/World/GeocodeServer/findAddressCandidates"
        geocode_params = {
            'SingleLine': f"{address}, Pinellas County, FL",
            'f': 'json',
            'outFields': '*'
        }
        
        geocode_response = session.get(search_url, params=geocode_params, timeout=15)
        geocode_data = geocode_response.json()
        
        if not geocode_data.get('candidates'):
            return {'success': False, 'error': 'Could not geocode address'}
        
        # Get coordinates
        location = geocode_data['candidates'][0]['location']
        x, y = location['x'], location['y']
        
        # Step 2: Query Zoning layer (Layer 1 - Zoning - Unincorporated)
        zoning_url = "https://egis.pinellas.gov/gis/rest/services/PublicWebGIS/Landuse_Zoning/MapServer/1/query"
        zoning_params = {
            'geometry': f"{x},{y}",
            'geometryType': 'esriGeometryPoint',
            'inSR': '4326',
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': 'ZONEDESC',
            'returnGeometry': 'false',
            'f': 'json'
        }
        
        zoning_response = session.get(zoning_url, params=zoning_params, timeout=15)
        zoning_data = zoning_response.json()
        
        zoning_code = ''
        zoning_desc = ''
        if zoning_data.get('features'):
            zoning_attrs = zoning_data['features'][0]['attributes']
            zoning_code = zoning_attrs.get('ZONEDESC', '')  # This returns the CODE
            # Look up the description from fetched or fallback dictionary
            zoning_desc = zoning_lookup.get(zoning_code, '')
        
        # Step 3: Query Future Land Use layer (Layer 0)
        flu_url = "https://egis.pinellas.gov/gis/rest/services/PublicWebGIS/Landuse_Zoning/MapServer/0/query"
        flu_params = {
            'geometry': f"{x},{y}",
            'geometryType': 'esriGeometryPoint',
            'inSR': '4326',
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': 'LANDUSECODE,LANDUSEDESC',
            'returnGeometry': 'false',
            'f': 'json'
        }
        
        flu_response = session.get(flu_url, params=flu_params, timeout=15)
        flu_data = flu_response.json()
        
        flu_code = ''
        flu_desc = ''
        if flu_data.get('features'):
            flu_attrs = flu_data['features'][0]['attributes']
            # Get the code from either field (they both return the code)
            flu_code = flu_attrs.get('LANDUSECODE') or flu_attrs.get('LANDUSEDESC', '')
            # Look up the description from fetched or fallback dictionary
            flu_desc = flu_lookup.get(flu_code, '')
        
        return {
            'success': True,
            'zoning_code': zoning_code,
            'zoning_description': zoning_desc,
            'future_land_use': flu_code,
            'future_land_use_description': flu_desc
        }
        
    except Exception as e:
        return {'success': False, 'error': f'Unincorporated zoning lookup error: {str(e)}'}

def lookup_clearwater_zoning(address):
    """
    Lookup zoning and FLU for Clearwater properties.
    Uses Clearwater's own GIS services.
    
    Returns: dict with zoning_code, zoning_description, future_land_use, future_land_use_description
    """
    if not address:
        return {'success': False, 'error': 'Address required for zoning lookup'}
    
    session = get_resilient_session()
    
    try:
        # Step 1: Geocode the address
        search_url = "https://geocode.arcgis.com/arcgis/rest/services/World/GeocodeServer/findAddressCandidates"
        geocode_params = {
            'SingleLine': f"{address}, Clearwater, FL",
            'f': 'json',
            'outFields': '*'
        }
        
        geocode_response = session.get(search_url, params=geocode_params, timeout=15)
        geocode_data = geocode_response.json()
        
        if not geocode_data.get('candidates'):
            return {'success': False, 'error': 'Could not geocode address'}
        
        # Get coordinates
        location = geocode_data['candidates'][0]['location']
        x, y = location['x'], location['y']
        
        # Step 2: Query Zoning layer
        zoning_lookup = fetch_coded_values(
            "https://gis.myclearwater.com/arcgis/rest/services/ArcGISMapServices/Zoning_WGS84/MapServer/1",
            "ZONING"
        )
        zoning_url = "https://gis.myclearwater.com/arcgis/rest/services/ArcGISMapServices/Zoning_WGS84/MapServer/1/query"
        zoning_params = {
            'geometry': f"{x},{y}",
            'geometryType': 'esriGeometryPoint',
            'inSR': '4326',
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': 'ZONING,ZONING_DESC',
            'returnGeometry': 'false',
            'f': 'json'
        }
        
        zoning_response = session.get(zoning_url, params=zoning_params, timeout=15)
        zoning_data = zoning_response.json()
        
        zoning_code = ''
        zoning_desc = ''
        if zoning_data.get('features'):
            zoning_attrs = zoning_data['features'][0]['attributes']
            zoning_code = zoning_attrs.get('ZONING', '')
            zoning_desc = zoning_attrs.get('ZONING_DESC', '') or zoning_lookup.get(zoning_code, '')
        
        # Step 3: Query Future Land Use layer (Layer 0 confirmed)
        # Fetch coded values for FLU
        flu_lookup = fetch_coded_values(
            "https://gis.myclearwater.com/arcgis/rest/services/ArcGISMapServices/FLU_w_PPC_Colors_WGS84/MapServer/0",
            "LU"
        )
        
        flu_url = "https://gis.myclearwater.com/arcgis/rest/services/ArcGISMapServices/FLU_w_PPC_Colors_WGS84/MapServer/0/query"
        flu_params = {
            'geometry': f"{x},{y}",
            'geometryType': 'esriGeometryPoint',
            'inSR': '4326',
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': 'LU',  # The actual field name!
            'returnGeometry': 'false',
            'f': 'json'
        }
        
        flu_response = session.get(flu_url, params=flu_params, timeout=15)
        flu_data = flu_response.json()
        
        flu_code = ''
        flu_desc = ''
        if flu_data.get('features'):
            flu_attrs = flu_data['features'][0]['attributes']
            flu_code = flu_attrs.get('LU', '')
            # Look up description from coded values
            flu_desc = flu_lookup.get(flu_code, '')
        
        return {
            'success': True,
            'zoning_code': zoning_code,
            'zoning_description': zoning_desc,
            'future_land_use': flu_code,
            'future_land_use_description': flu_desc
        }
        
    except Exception as e:
        return {'success': False, 'error': f'Clearwater zoning lookup error: {str(e)}'}

def lookup_largo_zoning(address, parcel_data=None):
    """
    Lookup zoning and FLU for Largo properties.
    Largo uses Future Land Use classification instead of traditional zoning.
    Queries Largo's parcel layer for Countywide_Plan_Map_Category_1.
    
    Args:
        address: Property address
        parcel_data: Not used (kept for compatibility)
    
    Returns: dict with zoning_code, future_land_use, descriptions
    """
    if not address:
        return {'success': False, 'error': 'Address required'}
    
    session = get_resilient_session()
    
    try:
        # Step 1: Geocode the address
        search_url = "https://geocode.arcgis.com/arcgis/rest/services/World/GeocodeServer/findAddressCandidates"
        geocode_params = {
            'SingleLine': f"{address}, Largo, FL",
            'f': 'json',
            'outFields': '*'
        }
        
        geocode_response = session.get(search_url, params=geocode_params, timeout=15)
        geocode_data = geocode_response.json()
        
        if not geocode_data.get('candidates'):
            return {'success': False, 'error': 'Could not geocode address'}
        
        # Get coordinates
        location = geocode_data['candidates'][0]['location']
        x, y = location['x'], location['y']
        
        # Step 2: Query Largo parcel layer for Countywide Plan category
        parcel_url = "https://maps.largo.com/arcgis/rest/services/Largo_GIS_Viewer_Map/MapServer/247/query"
        parcel_params = {
            'geometry': f"{x},{y}",
            'geometryType': 'esriGeometryPoint',
            'inSR': '4326',
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': 'Countywide_Plan_Map_Category_1',
            'returnGeometry': 'false',
            'f': 'json'
        }
        
        parcel_response = session.get(parcel_url, params=parcel_params, timeout=15)
        parcel_data_result = parcel_response.json()
        
        flu_value = ''
        if parcel_data_result.get('features'):
            attrs = parcel_data_result['features'][0]['attributes']
            flu_value = attrs.get('Countywide_Plan_Map_Category_1', '')
        
        if flu_value:
            return {
                'success': True,
                'zoning_code': flu_value,  # Largo uses FLU as zoning
                'zoning_description': None,  # Already combined in flu_value
                'future_land_use': flu_value,  # Same as zoning for Largo
                'future_land_use_description': None  # Already combined
            }
        else:
            return {
                'success': False,
                'error': 'Could not retrieve Countywide Plan category from Largo parcel layer'
            }
        
    except Exception as e:
        return {'success': False, 'error': f'Largo lookup error: {str(e)}'}


def lookup_pinellas_consolidated_zoning(address, city_name=None):
    """
    Lookup zoning/FLU using Pinellas Planning Council consolidated service.
    This service contains zoning data for ALL Pinellas municipalities in one place.
    
    Args:
        address: Property address (e.g., "200 CENTRAL AVE")
        city_name: Optional city name to help geocoding accuracy
        
    Returns:
        dict with zoning_code, zoning_description, future_land_use, future_land_use_description
    """
    session = get_resilient_session()
    
    try:
        # Step 1: Geocode the address to get coordinates
        search_url = "https://geocode.arcgis.com/arcgis/rest/services/World/GeocodeServer/findAddressCandidates"
        
        # Build search string with city if provided
        if city_name and city_name not in ['Unincorporated Pinellas', 'Unincorporated']:
            search_address = f"{address}, {city_name}, FL"
        else:
            search_address = f"{address}, Pinellas County, FL"
        
        geocode_params = {
            'SingleLine': search_address,
            'f': 'json',
            'outFields': '*'
        }
        
        geocode_response = session.get(search_url, params=geocode_params, timeout=15)
        geocode_data = geocode_response.json()
        
        if not geocode_data.get('candidates'):
            return {'success': False, 'error': 'Could not geocode address'}
        
        # Get coordinates from first candidate
        location = geocode_data['candidates'][0]['location']
        x, y = location['x'], location['y']
        
        # Step 2: Query Pinellas Planning Council consolidated zoning service
        # This service has layers for ALL Pinellas cities
        # We'll query the service and it will return whichever city layer contains the point
        ppc_base_url = "https://egis.pinellas.gov/gis/rest/services/AGO/PPC_Data/MapServer"
        
        # Calculate mapExtent (small buffer around point for identify operation)
        # Buffer of ~0.001 degrees (approximately 100 meters)
        buffer = 0.001
        map_extent = f"{x-buffer},{y-buffer},{x+buffer},{y+buffer}"
        
        # Query all layers at once using the root MapServer/identify endpoint
        zoning_params = {
            'geometry': f"{x},{y}",
            'geometryType': 'esriGeometryPoint',
            # identify uses 'sr' (not inSR/outSR)
            'sr': '4326',
            'layers': 'all',
            'mapExtent': map_extent,
            'imageDisplay': '400,300,96',
            'tolerance': 2,
            'returnGeometry': 'false',
            'f': 'json'
        }
        
        zoning_response = session.get(f"{ppc_base_url}/identify", params=zoning_params, timeout=15)
        zoning_data = zoning_response.json()
        
        if zoning_data.get('results') and len(zoning_data['results']) > 0:
            # Step 2: Pick the zoning result (do not assume results[0] is zoning)
            best = None
            for r in zoning_data['results']:
                a = (r or {}).get('attributes') or {}
                z = a.get('ZONING')
                if z not in (None, ''):
                    best = r
                    break
            if best is None:
                best = zoning_data['results'][0]

            attrs = (best or {}).get('attributes', {}) or {}
            zoning_code = attrs.get('ZONING', '') or ''
            jurisdiction = attrs.get('JURISDICTION', '') or ''

            return {
                'success': True,
                'zoning_code': zoning_code,
                'zoning_description': f"{jurisdiction} - {zoning_code}" if jurisdiction else zoning_code,
                'future_land_use': '',
                'future_land_use_description': '',
                'jurisdiction': jurisdiction
            }
        else:
            return {'success': False, 'error': 'No zoning found at address location'}
            
    except Exception as e:
        return {'success': False, 'error': f'Pinellas consolidated zoning lookup error: {str(e)}'}

def lookup_pinellas_zoning(city_name, address, parcel_data=None):
    """
    Router function: Lookup zoning for Pinellas County.
    PRIMARY: Uses consolidated Pinellas Planning Council service (all cities)
    FALLBACK: City-specific functions for enhanced data (St. Pete FLU, Largo FLU)
    
    Args:
        city_name: City name (e.g., "St. Petersburg", "Clearwater", "Largo")
        address: Property address (e.g., "200 CENTRAL AVE")
        parcel_data: Optional parcel data from PCPAO lookup (used for Largo)
        
    Returns:
        dict with zoning_code, zoning_description, future_land_use, future_land_use_description
    """
    if not address:
        return {'success': False, 'error': 'Address required for zoning lookup'}
    
    # PRIMARY: Try consolidated service first (covers all Pinellas cities)
    result = lookup_pinellas_consolidated_zoning(address, city_name)
    
    if result.get('success'):
        # ENHANCEMENT: For cities with detailed FLU data, get that too
        if 'St. Petersburg' in city_name or 'St Petersburg' in city_name:
            # Get St. Pete's detailed FLU codes
            stpete_result = lookup_stpete_zoning(address)
            if stpete_result.get('success'):
                result['future_land_use'] = stpete_result.get('future_land_use')
                result['future_land_use_description'] = stpete_result.get('future_land_use_description')
        
        elif 'Largo' in city_name:
            # Get Largo's specialized FLU
            largo_result = lookup_largo_zoning(address, parcel_data)
            if largo_result.get('success'):
                result['future_land_use'] = largo_result.get('future_land_use')
                result['future_land_use_description'] = largo_result.get('future_land_use_description')
        
        elif is_unincorporated(city_name):
            # Get unincorporated FLU
            uninc_result = lookup_unincorporated_zoning(address)
            if uninc_result.get('success'):
                result['future_land_use'] = uninc_result.get('future_land_use')
                result['future_land_use_description'] = uninc_result.get('future_land_use_description')
        
        return result
    
    # FALLBACK: If consolidated fails, try city-specific methods
    if 'St. Petersburg' in city_name or 'St Petersburg' in city_name:
        return lookup_stpete_zoning(address)
    elif 'Clearwater' in city_name:
        return lookup_clearwater_zoning(address)
    elif 'Largo' in city_name:
        return lookup_largo_zoning(address, parcel_data)
    elif is_unincorporated(city_name):
        return lookup_unincorporated_zoning(address)
    else:
        # Consolidated failed and no fallback available
        return {
            'success': False,
            'error': f'Could not find zoning data for {city_name}'
        }


def lookup_stpete_zoning(address):
    """
    Lookup zoning for St. Petersburg properties using St. Pete GIS layers.
    (Renamed from original lookup_pinellas_zoning St. Pete section)
    
    Args:
        address: Property address (e.g., "200 CENTRAL AVE")
        
    Returns:
        dict with zoning_code, zoning_description, future_land_use, future_land_use_description
    """
    session = get_resilient_session()
    
    try:
        # Step 1: Geocode the address to get coordinates
        search_url = "https://geocode.arcgis.com/arcgis/rest/services/World/GeocodeServer/findAddressCandidates"
        geocode_params = {
            'SingleLine': f"{address}, St. Petersburg, FL",
            'f': 'json',
            'outFields': '*'
        }
        
        geocode_response = session.get(search_url, params=geocode_params, timeout=15)
        geocode_data = geocode_response.json()
        
        if not geocode_data.get('candidates'):
            return {'success': False, 'error': 'Could not geocode address'}
        
        # Get coordinates from first candidate
        location = geocode_data['candidates'][0]['location']
        x, y = location['x'], location['y']
        
        # Step 2: Query zoning layer with coordinates (spatial query)
        zoning_url = "https://egis.stpete.org/arcgis/rest/services/ServicesDSD/Zoning/MapServer/2/query"
        zoning_params = {
            'geometry': f"{x},{y}",
            'geometryType': 'esriGeometryPoint',
            'inSR': '4326',  # WGS84 from geocoder
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': 'ZONECLASS,ZONEDESC',
            'returnGeometry': 'false',
            'f': 'json'
        }
        
        zoning_response = session.get(zoning_url, params=zoning_params, timeout=15)
        zoning_data = zoning_response.json()
        
        if zoning_data.get('features'):
            attrs = zoning_data['features'][0]['attributes']
            zoning_code = attrs.get('ZONECLASS', '')
            zoning_desc = ZONING_DESCRIPTIONS.get(zoning_code, attrs.get('ZONEDESC', ''))
            
            # Step 3: Query Future Land Use layer with same coordinates
            flu_url = "https://egis.stpete.org/arcgis/rest/services/ServicesDSD/Zoning/MapServer/4/query"
            flu_params = {
                'geometry': f"{x},{y}",
                'geometryType': 'esriGeometryPoint',
                'inSR': '4326',
                'spatialRel': 'esriSpatialRelIntersects',
                'outFields': '*',
                'returnGeometry': 'false',
                'f': 'json'
            }
            
            flu_response = session.get(flu_url, params=flu_params, timeout=15)
            flu_data = flu_response.json()
            flu_code = ''
            flu_desc = ''
            if flu_data.get('features'):
                flu_attrs = flu_data['features'][0].get('attributes', {})
                flu_code = flu_attrs.get('LANDUSECODE', '')
                flu_desc = FLU_DESCRIPTIONS.get(flu_code, '')
            
            return {
                'success': True,
                'zoning_code': zoning_code,
                'zoning_description': zoning_desc,
                'future_land_use': flu_code,
                'future_land_use_description': flu_desc
            }
        else:
            return {'success': False, 'error': 'No zoning found at address location'}
            
    except Exception as e:
        return {'success': False, 'error': f'St. Pete zoning lookup error: {str(e)}'}
    
    # Other cities (not St. Petersburg)
    return {
        'success': True,
        'zoning_code': 'Contact City/County for zoning',
        'zoning_description': None,
        'future_land_use': None,
        'future_land_use_description': None,
        'note': 'City-specific zoning data not available via API'
    }


# ============================================================================
# PROPOSAL GENERATOR STREAMLIT APP (TOKENIZED TEMPLATE + TAB WORKFLOW)
# ============================================================================

from typing import Dict, Any, List, Optional
import os
import datetime
import pandas as pd
from openpyxl import load_workbook
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

# Paths to the project files
BASE_DIR = pathlib.Path(__file__).parent
ECONTRACT_PATH = str(BASE_DIR / "data" / "eContractMaster.xlsm")
PROPOSAL_TASK_TOOL_PATH = str(BASE_DIR / "data" / "Proposal Task Tool.xlsx")
TEMPLATE_DOCX_PATH = str(BASE_DIR / "assets" / "Template.docx")

# --------------------------
# Proposal state
# --------------------------

def init_proposal_state() -> None:
    if "proposal" not in st.session_state:
        st.session_state.proposal = {
            "intake": {
                "county": "Pinellas",
                "municipality": "",
                "jurisdiction_display": "",
                "parcel_id": "",
                "address": "",
                "city": "",
                "zip": "",
                "owner": "",
                "land_use": "",
                "site_area_acres": "",
                "site_area_sqft": "",
                # Zoning / FLU split fields for tokens
                "zoning_code": "",
                "zoning_description": "",
                "flu_code": "",
                "flu_description": "",
                # convenience full strings (optional)
                "zoning": "",
                "future_land_use": "",
            },
            "client": {
                "client_name": "",
                "client_address": "",
                "client_city_state_zip": "",
                "client_contact_name": "",
                "client_contact_title": "",
                "client_contact_email": "",
                "client_contact_phone": "",
                "entity_name": "",
                "entity_address": "",
            },
            "project": {
                "project_name": "",
                "project_location": "",
                "project_understanding": "",
                "project_description": "",
                "project_acreage": "",
                "project_parcel_folio": "",
                "project_flu": "",
                "project_zoning": "",
                "exclusions": [],
                "assumptions": "",
                "schedule_summary": "",
                "proposal_date": "",
                "signer_name": "",
                "signer_title": "",
            },
            "scope": {
                "selected_tasks": [],  # list[dict]
                "additional_services": "",
            },
            "permits": {
                "selected_permits": [],
            },
            "fees": {
                "fee_lines": [],  # list[dict]
                "fee_type": "",
                "fee_total": "",
                "fee_paragraph": "",
                "fee_estimate_note": "",
                "invoice_email_to": "",
                "invoice_email_cc": "",
            },
        }

def _g(d: Dict[str, Any], k: str) -> str:
    v = d.get(k, "")
    return "" if v is None else str(v)

def build_token_map(proposal: Dict[str, Any]) -> Dict[str, str]:
    """
    Default token mapping for Template_tokens.docx.
    Missing values map to "" (empty string).
    """
    intake = proposal.get("intake", {}) or {}
    client = proposal.get("client", {}) or {}
    project = proposal.get("project", {}) or {}
    fees = proposal.get("fees", {}) or {}

    city = _g(intake, "city")
    state = "FL"
    city_state = f"{city}, {state}" if city else ""

    zoning_full = _g(intake, "zoning") or (
        f"{_g(intake,'zoning_code')} - {_g(intake,'zoning_description')}".strip(" -")
    )
    flu_full = _g(intake, "future_land_use") or (
        f"{_g(intake,'flu_code')} - {_g(intake,'flu_description')}".strip(" -")
    )

    # Default proposal date to today's long format if not set
    proposal_date = _g(project, "proposal_date").strip()
    if not proposal_date:
        proposal_date = datetime.datetime.now().strftime("%B %-d, %Y") if os.name != "nt" else datetime.datetime.now().strftime("%B %#d, %Y")

    return {
        "{{PROPOSAL_DATE}}": proposal_date,
        "{{CLIENT_NAME}}": _g(client, "client_name"),
        "{{CLIENT_CONTACT_NAME}}": _g(client, "client_contact_name"),
        "{{ENTITY_NAME}}": _g(client, "entity_name"),
        "{{ENTITY_ADDRESS}}": _g(client, "entity_address"),
        "{{PROJECT_NAME}}": _g(project, "project_name"),
        "{{PROJECT_LOCATION}}": _g(project, "project_location") or _g(intake, "address"),
        "{{CITY_STATE}}": city_state,
        "{{PROJECT_UNDERSTANDING}}": _g(project, "project_understanding"),
        "{{ASSUMPTIONS}}": _g(project, "assumptions"),
        "{{SCHEDULE_SUMMARY}}": _g(project, "schedule_summary"),
        "{{FEE_PARAGRAPH}}": _g(fees, "fee_paragraph"),
        "{{INVOICE_EMAIL_TO}}": _g(fees, "invoice_email_to"),
        "{{INVOICE_EMAIL_CC}}": _g(fees, "invoice_email_cc"),
        "{{SIGNER_NAME}}": _g(project, "signer_name"),
        "{{SIGNER_TITLE}}": _g(project, "signer_title"),

        # Useful extra tokens (even if not yet placed everywhere)
        "{{PARCEL_ID}}": _g(intake, "parcel_id"),
        "{{SITE_ADDRESS}}": _g(intake, "address"),
        "{{OWNER_NAME}}": _g(intake, "owner"),
        "{{LAND_USE}}": _g(intake, "land_use"),
        "{{SITE_AREA_ACRES}}": _g(intake, "site_area_acres"),
        "{{SITE_AREA_SF}}": _g(intake, "site_area_sqft"),
        "{{COUNTY}}": _g(intake, "county"),
        "{{MUNICIPALITY}}": _g(intake, "municipality"),
        "{{JURISDICTION_DISPLAY}}": _g(intake, "jurisdiction_display"),
        "{{ZONING_CODE}}": _g(intake, "zoning_code"),
        "{{ZONING_DESCRIPTION}}": _g(intake, "zoning_description"),
        "{{ZONING_FULL}}": zoning_full,
        "{{FLU_CODE}}": _g(intake, "flu_code"),
        "{{FLU_DESCRIPTION}}": _g(intake, "flu_description"),
        "{{FLU_FULL}}": flu_full,
        "{{FEE_TYPE}}": _g(fees, "fee_type"),
        "{{FEE_TOTAL}}": _g(fees, "fee_total"),
        "{{FEE_NOTE}}": _g(fees, "fee_estimate_note"),
    }

# --------------------------
# Task database load
# --------------------------

def is_relevant_for_jurisdiction(municipality: Optional[str], current_county: Optional[str] = None) -> bool:
    """
    Check if a task is relevant for the current jurisdiction.
    Generic tasks (no municipality) are always relevant.
    If a county is specified, filter out tasks from other regions.
    """
    if not municipality:
        return True

    # If we have no county context, include all tasks
    if not current_county:
        return True

    # Now check if task municipality matches the current jurisdiction context
    muni_upper = str(municipality).upper()
    county_upper = str(current_county).upper()

    # Check if the task's municipality is in the current county or matches it
    if county_upper in muni_upper:
        return True

    # For backward compatibility: filter out Orlando area if we're in Tampa Bay area
    orlando_indicators = ["ORLANDO", "ORANGE COUNTY", "WINTER GARDEN", "OVIEDO", "APOPKA",
                          "WINTER PARK", "CASSELBERRY", "LAKE COUNTY", "OSCEOLA COUNTY"]

    pinellas_indicators = ["PINELLAS", "ST. PETERSBURG", "CLEARWATER", "LARGO", "TAMPA"]

    is_orlando_task = any(indicator in muni_upper for indicator in orlando_indicators)
    is_pinellas_context = any(indicator in county_upper for indicator in pinellas_indicators)

    # If we're in Pinellas context, exclude Orlando tasks
    if is_pinellas_context and is_orlando_task:
        return False

    return True

def stable_task_id(*parts: str) -> str:
    import hashlib
    raw = "||".join([p or "" for p in parts]).encode("utf-8")
    return hashlib.sha1(raw).hexdigest()[:12]

@st.cache_resource(show_spinner=False)
def load_task_database() -> Dict[str, Any]:
    db: Dict[str, Any] = {
        "econtract_tasks": {},
        "category_tasks": {},
        "meta": {"loaded_at": datetime.datetime.now().isoformat(timespec="seconds")},
    }

    # eContract TaskDescriptions (available for later matching)
    try:
        wb = load_workbook(ECONTRACT_PATH, data_only=True)
        if "TaskDescriptions" in wb.sheetnames:
            sh = wb["TaskDescriptions"]
            current_task = None
            current_desc_lines: List[str] = []
            for row in sh.iter_rows(min_row=2, values_only=True):
                desc_text = str(row[0]).strip() if row and row[0] else None
                if not desc_text:
                    continue
                if desc_text.startswith("Task "):
                    if current_task is not None:
                        db["econtract_tasks"][current_task]["description"] = "\n".join(current_desc_lines).strip()
                    parts = desc_text.split(" - ", 1)
                    try:
                        task_no = int(parts[0].replace("Task", "").strip())
                    except Exception:
                        task_no = None
                    if task_no is not None:
                        task_name = parts[1].strip() if len(parts) > 1 else f"Task {task_no}"
                        db["econtract_tasks"][task_no] = {"name": task_name, "description": ""}
                        current_task = task_no
                        current_desc_lines = []
                    else:
                        current_task = None
                        current_desc_lines = []
                else:
                    if current_task is not None:
                        current_desc_lines.append(desc_text)
            if current_task is not None:
                db["econtract_tasks"][current_task]["description"] = "\n".join(current_desc_lines).strip()
    except Exception as e:
        db["meta"]["econtract_error"] = str(e)

    # Proposal Task Tool sheets
    try:
        wb2 = load_workbook(PROPOSAL_TASK_TOOL_PATH, data_only=True)
        for sheet_name in wb2.sheetnames:
            sh2 = wb2[sheet_name]
            if sheet_name.strip().lower() in {"readme", "instructions"}:
                continue
            tasks: List[Dict[str, Any]] = []
            for row in sh2.iter_rows(min_row=2, values_only=True):
                municipality = row[0] if len(row) > 0 else None
                task_name = row[1] if len(row) > 1 else None
                description = row[2] if len(row) > 2 else None
                if not task_name or not description:
                    continue
                # Will be filtered later based on actual county context
                # if not is_relevant_for_jurisdiction(municipality):
                #     continue
                municipality_str = str(municipality).strip() if municipality else ""
                task_name_str = str(task_name).strip()
                description_str = str(description).strip()
                task_id = stable_task_id(sheet_name, municipality_str, task_name_str)
                tasks.append({
                    "task_id": task_id,
                    "category": sheet_name,
                    "municipality": municipality_str,
                    "task_name": task_name_str,
                    "description": description_str,
                    "is_generic": (municipality_str == ""),
                    "source": "Proposal Task Tool",
                })
            if tasks:
                db["category_tasks"][sheet_name] = tasks
    except Exception as e:
        db["meta"]["proposal_tool_error"] = str(e)

    return db

# --------------------------
# DOCX helpers (token replacement + anchors)
# --------------------------

def _replace_text_in_paragraph(p: Paragraph, token_map: Dict[str, str]) -> None:
    full = "".join(run.text for run in p.runs)
    if not full:
        return
    new = full
    for k, v in token_map.items():
        if k in new:
            new = new.replace(k, v)
    if new != full:
        for r in p.runs:
            r.text = ""
        if p.runs:
            p.runs[0].text = new
        else:
            p.add_run(new)

def replace_tokens_everywhere(doc: Document, token_map: Dict[str, str]) -> None:
    for p in doc.paragraphs:
        _replace_text_in_paragraph(p, token_map)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_text_in_paragraph(p, token_map)

def delete_paragraph(p: Paragraph) -> None:
    p._element.getparent().remove(p._element)

def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para

def replace_anchor_with_block(doc: Document, anchor_token: str, lines: List[str], bullet: bool = False) -> None:
    """
    Finds a paragraph whose text contains anchor_token.
    Replaces that paragraph with the provided lines (as bullets or plain).
    """
    for p in list(doc.paragraphs):
        if anchor_token in p.text:
            # clear anchor paragraph text
            p.text = ""
            if not lines:
                delete_paragraph(p)
                return
            # first line in existing paragraph
            if bullet:
                p.style = doc.styles["List Bullet"] if "List Bullet" in doc.styles else p.style
                p.add_run(lines[0])
            else:
                p.add_run(lines[0])

            prev = p
            for line in lines[1:]:
                np = insert_paragraph_after(prev, "")
                if bullet:
                    np.style = doc.styles["List Bullet"] if "List Bullet" in doc.styles else np.style
                    np.add_run(line)
                else:
                    np.add_run(line)
                prev = np
            return

def fill_fee_table_anchor(doc: Document, anchor_token: str, fee_lines: List[Dict[str, Any]]) -> None:
    """
    Finds the first table cell containing anchor_token and fills that table as a fee summary.
    Assumes row 0 is header; clears all rows below header then appends one row per fee line.
    """
    for tbl in doc.tables:
        found = False
        for r in tbl.rows:
            for c in r.cells:
                if anchor_token in c.text:
                    found = True
                    c.text = ""  # clear token
                    break
            if found:
                break
        if found:
            # Keep header row only
            while len(tbl.rows) > 1:
                tbl._tbl.remove(tbl.rows[1]._tr)

            for fl in fee_lines:
                row = tbl.add_row().cells
                # Try to match the template header meaningfully; if it has 3 cols: Task | Type | Fee
                if len(row) >= 1:
                    row[0].text = str(fl.get("task_name",""))
                if len(row) >= 2:
                    row[1].text = str(fl.get("fee_type",""))
                if len(row) >= 3:
                    row[2].text = str(fl.get("fee_amount",""))
            return

def generate_proposal_docx(output_path: str, proposal: Dict[str, Any]) -> str:
    doc = Document(TEMPLATE_DOCX_PATH)

    # 1) Replace token text
    token_map = build_token_map(proposal)
    replace_tokens_everywhere(doc, token_map)

    # 2) Anchors (scope / permits / additional services / fee table)
    scope_tasks = proposal.get("scope", {}).get("selected_tasks", []) or []
    scope_lines: List[str] = []
    for t in scope_tasks:
        name = t.get("task_name","").strip()
        muni = t.get("municipality","").strip()
        desc = (t.get("description","") or "").strip()
        if muni:
            scope_lines.append(f"{name} ({muni})")
        else:
            scope_lines.append(name)
        if desc:
            # Keep description as following line (indented look without custom styles)
            scope_lines.append(desc)

    replace_anchor_with_block(doc, "{{INSERT_SCOPE_TASKS}}", scope_lines, bullet=False)

    permits = proposal.get("permits", {}).get("selected_permits", []) or []
    replace_anchor_with_block(doc, "{{INSERT_PERMITS}}", permits, bullet=True)

    add_services_text = (proposal.get("scope", {}) or {}).get("additional_services","") or ""
    add_lines = [ln.strip() for ln in add_services_text.splitlines() if ln.strip()]
    replace_anchor_with_block(doc, "{{INSERT_ADDITIONAL_SERVICES}}", add_lines, bullet=True)

    fee_lines = proposal.get("fees", {}).get("fee_lines", []) or []
    fill_fee_table_anchor(doc, "{{INSERT_FEE_TABLE}}", fee_lines)

    doc.save(output_path)
    return output_path

# ============================================================================
# STREAMLIT UI (6 TAB LAYOUT)
# ============================================================================

init_proposal_state()
db = load_task_database()

st.set_page_config(page_title="Proposal Generator", page_icon="🧾", layout="wide")
st.title("🧾 Proposal Generator (Tokenized)")
st.caption("Intake → Understanding → Scope → Permits → Fees → Review & Generate")


st.markdown(
    """
<style>
body, .stApp {
  background-color: #c7c7c7;
  color: #000000;
}
.stTextInput input,
.stTextArea textarea,
.stNumberInput input,
.stDateInput input {
  background-color: #ffffff !important;
  color: #000000 !important;
  border: 2px solid #000000 !important;
  box-shadow: none !important;
}
.stTextInput input:focus,
.stTextArea textarea:focus,
.stNumberInput input:focus,
.stDateInput input:focus {
  outline: none !important;
  border: 2px solid #000000 !important;
  box-shadow: none !important;
}
.stSelectbox div[data-baseweb="select"] > div,
.stMultiSelect div[data-baseweb="select"] > div {
  background-color: #ffffff !important;
  color: #000000 !important;
  border: 2px solid #000000 !important;
  box-shadow: none !important;
}
.stSelectbox div[data-baseweb="select"] input,
.stMultiSelect div[data-baseweb="select"] input {
  border: 0 !important;
  box-shadow: none !important;
  outline: none !important;
  caret-color: transparent !important;
  background: transparent !important;
}
.stSelectbox div[data-baseweb="select"] input:focus,
.stSelectbox div[data-baseweb="select"] input:active,
.stMultiSelect div[data-baseweb="select"] input:focus,
.stMultiSelect div[data-baseweb="select"] input:active {
  border: 0 !important;
  box-shadow: none !important;
  outline: none !important;
}
.stSelectbox div[data-baseweb="select"] div[role="listbox"],
.stMultiSelect div[data-baseweb="select"] div[role="listbox"] {
  border: 0 !important;
  box-shadow: none !important;
  outline: none !important;
}
.stSelectbox div[data-baseweb="select"] [data-baseweb="tag"],
.stSelectbox div[data-baseweb="select"] [data-baseweb="input"],
.stMultiSelect div[data-baseweb="select"] [data-baseweb="tag"],
.stMultiSelect div[data-baseweb="select"] [data-baseweb="input"] {
  border: 0 !important;
  box-shadow: none !important;
  outline: none !important;
}
.stSelectbox div[data-baseweb="select"] input::selection,
.stMultiSelect div[data-baseweb="select"] input::selection {
  background: transparent !important;
}
.stSelectbox div[data-baseweb="select"] input::placeholder,
.stMultiSelect div[data-baseweb="select"] input::placeholder {
  color: #000000 !important;
}
.stApp * {
  caret-color: #000000;
}
.stTextInput input:disabled,
.stTextArea textarea:disabled,
.stNumberInput input:disabled,
.stDateInput input:disabled,
.stSelectbox div[data-baseweb="select"] input:disabled,
.stSelectbox div[data-baseweb="select"] div[aria-disabled="true"],
.stMultiSelect div[data-baseweb="select"] div[aria-disabled="true"] {
  background-color: #ffffff !important;
  color: #444444 !important;
  opacity: 1 !important;
}
</style>
""",
    unsafe_allow_html=True,
)

tabs = st.tabs([
    "📍 Intake",
    "🧠 Project Understanding",
    "✅ Scope",
    "🗺️ Permits",
    "🧾 Fees",
    "🧾 Review & Generate",
])

# -------------------------
# TAB 1: Intake (Lookup)
# -------------------------
with tabs[0]:
    st.subheader("📍 Intake (Lookup)")
    left, right = st.columns([1, 1])

    with left:
        st.markdown("**Property Lookup**")

        intake = st.session_state.proposal["intake"]
        parcel_col, county_col = st.columns([2, 1])
        with parcel_col:
            parcel_id_input = st.text_input(
                "Parcel ID",
                value=intake.get("parcel_id", ""),
                placeholder="e.g., 19-31-17-73166-001-0010",
                help="Parcel ID with dashes",
            )
        with county_col:
            county_options = ["Pinellas", "Hillsborough", "Pasco"]
            current_county = intake.get("county", "Pinellas") or "Pinellas"
            try:
                county_index = county_options.index(current_county)
            except ValueError:
                county_index = 0
            county_input = st.selectbox("County", options=county_options, index=county_index)
            intake["county"] = county_input

        if st.button("Lookup Property Data", type="primary", use_container_width=True, key="lookup_property"):
            if not parcel_id_input:
                st.error("Please enter a parcel ID.")
            else:
                is_valid, error_msg = validate_parcel_id(parcel_id_input)
                if not is_valid:
                    st.error(f"{error_msg}")
                elif county_input != "Pinellas":
                    intake["county"] = county_input
                    intake["parcel_id"] = parcel_id_input
                    st.error("Property lookup is only implemented for Pinellas County right now.")
                else:
                    # Step 1: Fetch property info
                    with st.spinner("Fetching property data from PCPAO API..."):
                        result = scrape_pinellas_property(parcel_id_input)

                    if result.get("success"):
                        intake["county"] = county_input
                        intake["parcel_id"] = parcel_id_input
                        intake["address"] = result.get("address", "") or ""
                        intake["city"] = expand_city_name(result.get("city", "") or "")
                        intake["zip"] = result.get("zip", "") or ""
                        intake["owner"] = result.get("owner", "") or ""
                        intake["land_use"] = result.get("land_use", "") or ""
                        intake["site_area_sqft"] = result.get("site_area_sqft", "") or ""
                        intake["site_area_acres"] = result.get("site_area_acres", "") or ""
                        intake["municipality"] = intake["city"]
                        intake["jurisdiction_display"] = intake["city"]

                        st.success("Property data retrieved.")
                        st.rerun()
                    else:
                        st.error(f"{result.get('error','Lookup failed')}")

        if st.button("Lookup Zoning / Future Land Use", use_container_width=True, key="lookup_zoning"):
            intake["county"] = county_input
            city = expand_city_name(intake.get("city", "") or "")
            address = intake.get("address", "")
            if not address or not city:
                st.error("Address and city are required to lookup zoning/FLU.")
            elif county_input != "Pinellas":
                st.error("Zoning/FLU lookup is only implemented for Pinellas County right now.")
            else:
                with st.spinner(f"Fetching zoning/FLU for {address}..."):
                    zoning_result = lookup_pinellas_zoning(city, address, parcel_data={'parcel_id': intake.get('parcel_id','')})

                if zoning_result.get("success"):
                    intake["zoning_code"] = zoning_result.get("zoning_code") or ""
                    intake["zoning_description"] = zoning_result.get("zoning_description") or ""
                    intake["flu_code"] = zoning_result.get("future_land_use") or ""
                    intake["flu_description"] = zoning_result.get("future_land_use_description") or ""

                    if intake["zoning_code"] and intake["zoning_description"]:
                        intake["zoning"] = f"{intake['zoning_code']} - {intake['zoning_description']}"
                    else:
                        intake["zoning"] = intake["zoning_code"] or ""

                    if intake["flu_code"] and intake["flu_description"]:
                        intake["future_land_use"] = f"{intake['flu_code']} - {intake['flu_description']}"
                    else:
                        intake["future_land_use"] = intake["flu_code"] or ""

                    # Treat "success but empty results" as a failure for user visibility.
                    missing = []
                    if not intake["zoning_code"]:
                        missing.append("zoning")
                    if not intake["flu_code"]:
                        missing.append("future land use")

                    if missing:
                        st.warning(f"Lookup completed but returned no {', '.join(missing)} result(s).")
                        dbg = zoning_result.get("_debug") or zoning_result
                        if dbg:
                            with st.expander("Debug details"):
                                st.json(dbg)
                    else:
                        st.success("Zoning/FLU data retrieved.")
                        st.rerun()
                else:
                    st.error(f"Zoning/FLU lookup failed: {zoning_result.get('error','Unknown error')}")
                    dbg = zoning_result.get("_debug") or zoning_result
                    if dbg:
                        with st.expander("Debug details"):
                            st.json(dbg)



        st.markdown("**Lookup Summary (Auto-fills tokens)**")
        st.text_input("County", value=intake.get("county",""), disabled=True)
        st.text_input("City", value=intake.get("city",""), disabled=True)
        st.text_input("Address", value=intake.get("address",""), disabled=True)
        st.text_input("Owner", value=intake.get("owner",""), disabled=True)
        st.text_input("Land Use", value=intake.get("land_use",""), disabled=True)
        st.text_input("Zoning (full)", value=intake.get("zoning",""), disabled=True)
        st.text_input("Future Land Use (full)", value=intake.get("future_land_use",""), disabled=True)
        st.text_input("Site Area (acres)", value=intake.get("site_area_acres",""), disabled=True)
        st.text_input("Site Area (sf)", value=intake.get("site_area_sqft",""), disabled=True)

    with right:
        project = st.session_state.proposal["project"]
        default_loc = project.get("project_location","") or intake.get("address","")

        st.markdown("**Project (Tokens)**")
        project["project_name"] = st.text_input("Project Name", value=project.get("project_name",""))
        project["project_location"] = st.text_input("Project Location / Address", value=default_loc)
        project["proposal_date"] = st.text_input("Proposal Date (optional)", value=project.get("proposal_date",""), help="Leave blank to auto-use today's date.")

        st.markdown("**Client / Entity (Tokens)**")
        client = st.session_state.proposal["client"]
        client["client_name"] = st.text_input("Client Name", value=client.get("client_name",""))
        client["client_contact_name"] = st.text_input("Client Contact Name", value=client.get("client_contact_name",""))
        client["entity_name"] = st.text_input("Client Legal Entity (Sunbiz)", value=client.get("entity_name",""))
        client["entity_address"] = st.text_area("Entity Address", value=client.get("entity_address",""), height=90)


# -------------------------
# TAB 2: Project Understanding
# -------------------------
with tabs[1]:
    st.subheader("🧠 Project Understanding")
    project = st.session_state.proposal["project"]
    intake = st.session_state.proposal["intake"]

    st.markdown("---")
    st.markdown("**Project Understanding Template**")
    st.caption("Fill in the details below. This will generate the project understanding paragraph.")

    project["project_description"] = st.text_area(
        "Project Description",
        value=project.get("project_description",""),
        height=100,
        help="What the client plans to develop (e.g., 'a 50-unit residential subdivision')"
    )

    st.markdown("---")
    st.markdown("**Tasks NOT Included in Scope**")
    st.caption("Select items that are explicitly excluded from the project scope.")

    common_exclusions = [
        "Geotechnical investigation",
        "Topographic survey",
        "Boundary survey",
        "Environmental assessment (Phase I/II)",
        "Traffic impact analysis",
        "Lighting design",
        "Landscape architecture",
        "Architectural services",
        "Structural engineering",
        "MEP (Mechanical, Electrical, Plumbing) engineering",
        "Permitting fees",
        "Application fees",
        "Construction administration",
        "Construction observation",
        "As-built survey",
    ]

    if "exclusions" not in project:
        project["exclusions"] = []

    selected_exclusions = project.get("exclusions", [])

    cols = st.columns(3)
    updated_exclusions = []

    for idx, exclusion in enumerate(common_exclusions):
        col_idx = idx % 3
        with cols[col_idx]:
            if st.checkbox(exclusion, value=(exclusion in selected_exclusions), key=f"exclusion_{exclusion}"):
                updated_exclusions.append(exclusion)

    project["exclusions"] = updated_exclusions

    # Auto-generate project understanding paragraph
    if st.button("Generate Project Understanding Paragraph"):
        proj_desc = project.get("project_description", "XX")
        proj_loc = project.get("project_location", "") or intake.get("address", "project general location or address")
        proj_parcel = intake.get("parcel_id", "")
        proj_acres = intake.get("site_area_acres", "X")
        proj_flu = intake.get("future_land_use", "") or intake.get("flu_code", "X")
        proj_zoning = intake.get("zoning", "") or intake.get("zoning_code", "X")

        generated_text = f"Kimley-Horn understands that the Client plans to develop {proj_desc} on the property located at {proj_loc}"
        if proj_parcel:
            generated_text += f" (Parcel #{proj_parcel})"
        generated_text += f". The {proj_acres}-acre parcel has a Future Land Use designation of {proj_flu} and is Zoned {proj_zoning}."
        if project.get("exclusions"):
            excluded = "; ".join(project["exclusions"])
            generated_text += f" The following items are excluded from scope: {excluded}."

        project["project_understanding"] = generated_text
        st.success("Project understanding paragraph generated!")
        st.rerun()

    project["project_understanding"] = st.text_area(
        "Project Understanding (editable)",
        value=project.get("project_understanding",""),
        height=180,
        help="This goes into the template. You can edit the auto-generated text or write your own."
    )

    st.markdown("---")
    project["assumptions"] = st.text_area("Assumptions / Notes", value=project.get("assumptions",""), height=160)
    project["schedule_summary"] = st.text_area("Schedule Summary", value=project.get("schedule_summary",""), height=120)

# -------------------------
# TAB 3: Scope
# -------------------------
with tabs[2]:
    st.subheader("✅ Scope")
    st.caption("Select tasks from Proposal Task Tool.xlsx. These feed {{INSERT_SCOPE_TASKS}}.")

    # Cache clear button
    if st.button("🔄 Reload Task Database", help="Clear cache and reload tasks from Excel files"):
        load_task_database.clear()
        st.rerun()

    categories = sorted(db.get("category_tasks", {}).keys())
    if not categories:
        st.error("No tasks loaded from Proposal Task Tool.xlsx. Check file path and sheet formats.")

        # Debug info
        with st.expander("Debug Info"):
            st.write("Expected path:", PROPOSAL_TASK_TOOL_PATH)
            st.write("File exists:", pathlib.Path(PROPOSAL_TASK_TOOL_PATH).exists())
            st.write("Database meta:", db.get("meta", {}))
            if "proposal_tool_error" in db.get("meta", {}):
                st.error(f"Error loading tasks: {db['meta']['proposal_tool_error']}")
    else:
        desired_order = [
            "Due Diligence",
            "Comp Plan zoning",
            "Pre-app",
            "SD-DD",
            "CDs",
            "Misc Tasks",
            "Meetings",
            "Permitting",
            "Utility",
        ]

        categories_map = {c.lower(): c for c in categories}
        ordered_categories = [
            categories_map[name.lower()]
            for name in desired_order
            if name.lower() in categories_map
        ]

        selected_ids = {t["task_id"] for t in st.session_state.proposal["scope"]["selected_tasks"]}
        new_selected = []

        for cat in ordered_categories:
            st.markdown(f"**{cat}**")
            tasks = db.get("category_tasks", {}).get(cat, [])
            tasks = sorted(tasks, key=lambda t: t.get("task_name", ""))
            if not tasks:
                st.info("No tasks available in this category.")
            for t in tasks:
                label = t.get("task_name", "").strip()
                muni = (t.get("municipality") or "").strip()
                if muni:
                    label = f"{label} ({muni})"
                checked = st.checkbox(
                    label,
                    value=(t.get("task_id") in selected_ids),
                    key=f"scope_task_{cat}_{t.get('task_id')}",
                )
                if checked:
                    new_selected.append(t)
            st.markdown("---")

        st.session_state.proposal["scope"]["selected_tasks"] = new_selected

        if st.button("🗑️ Clear Selected Tasks"):
            st.session_state.proposal["scope"]["selected_tasks"] = []
            st.session_state.proposal["fees"]["fee_lines"] = []
            st.rerun()

    st.markdown("---")
    st.markdown("**Additional Services** (feeds {{INSERT_ADDITIONAL_SERVICES}})")
    st.session_state.proposal["scope"]["additional_services"] = st.text_area(
        "Optional additional services to include",
        value=st.session_state.proposal["scope"].get("additional_services",""),
        height=120
    )

# -------------------------
# TAB 4: Permits
# -------------------------
with tabs[3]:
    st.subheader("🗺️ Permits")
    st.caption("Selected permits feed {{INSERT_PERMITS}}.")

    common_permits = [
        "AHJ Site Development / Civil Plan Approval",
        "Water Provider Approval",
        "Sewer Provider Approval",
        "SWFWMD Environmental Resource Permit (ERP)",
        "FDEP Potable Water / Wastewater (if applicable)",
        "FDOT Driveway Connection Permit (if applicable)",
        "FDOT Drainage Connection Permit (if applicable)",
        "FEMA (if applicable)",
    ]

    # Initialize selected permits if not exists
    if "selected_permits" not in st.session_state.proposal["permits"]:
        st.session_state.proposal["permits"]["selected_permits"] = []

    selected_permits = st.session_state.proposal["permits"]["selected_permits"]

    st.markdown("**Select Required Permits:**")

    # Create checkboxes for each permit
    updated_permits = []
    for permit in common_permits:
        if st.checkbox(permit, value=(permit in selected_permits), key=f"permit_{permit}"):
            updated_permits.append(permit)

    st.session_state.proposal["permits"]["selected_permits"] = updated_permits

# -------------------------
# TAB 5: Fees
# -------------------------
with tabs[4]:
    st.subheader("🧾 Fees")
    st.caption("Fee lines populate the fee table anchor in Template_tokens.docx.")

    fees = st.session_state.proposal["fees"]

    fees["fee_type"] = st.selectbox(
        "Overall Fee Type (token: {{FEE_TYPE}})",
        options=["", "Lump Sum", "Hourly", "Hourly, Not-to-Exceed"],
        index=["", "Lump Sum", "Hourly", "Hourly, Not-to-Exceed"].index(fees.get("fee_type","") if fees.get("fee_type","") in ["", "Lump Sum", "Hourly", "Hourly, Not-to-Exceed"] else "")
    )
    fees["fee_total"] = st.text_input("Overall Fee Total (token: {{FEE_TOTAL}})", value=fees.get("fee_total",""))

    # Default fee paragraph builder (editable)
    default_fee_para = fees.get("fee_paragraph","")
    if not default_fee_para:
        if fees["fee_type"] == "Lump Sum" and fees["fee_total"]:
            default_fee_para = f"Kimley-Horn is pleased to provide the services described herein for a lump sum fee of {fees['fee_total']}."
        elif fees["fee_type"] == "Hourly":
            default_fee_para = "Kimley-Horn will provide the services described herein on an hourly basis in accordance with the attached or applicable rate schedule."
        elif fees["fee_type"] == "Hourly, Not-to-Exceed" and fees["fee_total"]:
            default_fee_para = f"Kimley-Horn will provide the services described herein on an hourly basis with a total fee not-to-exceed {fees['fee_total']}."
    fees["fee_paragraph"] = st.text_area("Fee Paragraph (token: {{FEE_PARAGRAPH}})", value=default_fee_para, height=130)

    fees["invoice_email_to"] = st.text_input("Invoice Email To (token: {{INVOICE_EMAIL_TO}})", value=fees.get("invoice_email_to",""))
    fees["invoice_email_cc"] = st.text_input("Invoice Email CC (token: {{INVOICE_EMAIL_CC}})", value=fees.get("invoice_email_cc",""))

    fees["fee_estimate_note"] = st.text_area("Fee note (token: {{FEE_NOTE}})", value=fees.get("fee_estimate_note",""), height=100)

    sel = st.session_state.proposal["scope"]["selected_tasks"]
    if not sel:
        st.info("Select tasks in the Scope tab first.")
    else:
        fee_lines = fees.get("fee_lines", [])
        fee_by_id = {fl.get("task_id"): fl for fl in fee_lines if fl.get("task_id")}

        normalized = []
        for t in sel:
            tid = t["task_id"]
            if tid in fee_by_id:
                fl = fee_by_id[tid]
            else:
                fl = {"task_id": tid, "task_name": t.get("task_name",""), "fee_type": fees.get("fee_type","Lump Sum") or "Lump Sum", "fee_amount": ""}
            normalized.append(fl)

        fee_df = pd.DataFrame(normalized)
        edited = st.data_editor(
            fee_df,
            use_container_width=True,
            num_rows="fixed",
            column_config={
                "fee_type": st.column_config.SelectboxColumn(
                    "Fee Type",
                    options=["Lump Sum", "Hourly", "Hourly, Not-to-Exceed"],
                    required=True,
                )
            }
        )
        fees["fee_lines"] = edited.to_dict(orient="records")

# -------------------------
# TAB 6: Review & Generate
# -------------------------
with tabs[5]:
    st.subheader("🧾 Review & Generate")

    p = st.session_state.proposal
    intake = p["intake"]
    client = p["client"]
    project = p["project"]
    scope_tasks = p["scope"]["selected_tasks"]

    missing = []
    if not client.get("client_name"):
        missing.append("Client name")
    if not client.get("client_contact_name"):
        missing.append("Client contact name (salutation token)")
    if not client.get("entity_name"):
        missing.append("Client legal entity")
    if not project.get("project_name"):
        missing.append("Project name")
    if not (project.get("project_location") or intake.get("address")):
        missing.append("Project location/address")
    if not scope_tasks:
        missing.append("At least one scope task (Scope tab)")

    if missing:
        st.warning("Missing required items: " + ", ".join(missing))

    st.markdown("### Token Preview (key ones)")
    token_map = build_token_map(p)
    preview_keys = [
        "{{PROPOSAL_DATE}}","{{CLIENT_NAME}}","{{CLIENT_CONTACT_NAME}}","{{ENTITY_NAME}}",
        "{{PROJECT_NAME}}","{{PROJECT_LOCATION}}","{{CITY_STATE}}","{{ZONING_FULL}}","{{FLU_FULL}}",
        "{{SITE_AREA_ACRES}}","{{PARCEL_ID}}","{{FEE_PARAGRAPH}}"
    ]
    st.json({k: token_map.get(k,"") for k in preview_keys})

    st.markdown("### Output")
    out_name = st.text_input("Output filename", value="Proposal_Draft_Tokens.docx")
    out_path = str(BASE_DIR / "output" / out_name)

    can_generate = (len(missing) == 0)

    if st.button("Generate DOCX (Token Template)", type="primary", disabled=not can_generate):
        try:
            generated = generate_proposal_docx(out_path, p)
            with open(generated, "rb") as f:
                st.download_button(
                    "⬇️ Download Proposal DOCX",
                    data=f,
                    file_name=os.path.basename(generated),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            st.success("Generated.")
        except Exception as e:
            st.error(f"Generation failed: {e}")
